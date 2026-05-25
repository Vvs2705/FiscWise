"""Document Parsing Service for FiscWise

Handles parsing of different document formats (PDF, Excel, Word).
Extracts text and structured data from uploaded documents.
"""

from typing import Optional, Dict, Any
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


async def parse_document(file_path: str, file_name: str) -> tuple[Optional[Dict[str, Any]], Optional[str]]:
    """
    Parse document based on file type.

    Args:
        file_path: Path to uploaded file
        file_name: Original file name

    Returns:
        Tuple of (parsed_data, error_message)
        parsed_data is a dict with extracted content
        error_message is set if parsing failed
    """
    try:
        file_ext = Path(file_name).suffix.lower()

        if file_ext == ".pdf":
            return await _parse_pdf(file_path)
        elif file_ext in [".xlsx", ".xls"]:
            return await _parse_excel(file_path)
        elif file_ext in [".docx", ".doc"]:
            return await _parse_word(file_path)
        elif file_ext == ".txt":
            return await _parse_text(file_path)
        else:
            return None, f"Unsupported file format: {file_ext}"

    except Exception as e:
        logger.error(f"Error parsing document: {str(e)}")
        return None, str(e)


async def _parse_pdf(file_path: str) -> tuple[Optional[Dict[str, Any]], Optional[str]]:
    """Parse PDF file and extract text."""
    try:
        import pdfplumber

        data = {
            "type": "pdf",
            "pages": [],
            "total_pages": 0,
            "full_text": ""
        }

        with pdfplumber.open(file_path) as pdf:
            data["total_pages"] = len(pdf.pages)

            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text() or ""
                data["pages"].append({
                    "page": page_num,
                    "text": text[:1000]  # First 1000 chars per page
                })
                data["full_text"] += text + "\n"

        # Truncate full text to first 5000 chars
        data["full_text"] = data["full_text"][:5000]
        return data, None

    except ImportError:
        return None, "pdfplumber library not installed"
    except Exception as e:
        return None, f"PDF parsing error: {str(e)}"


async def _parse_excel(file_path: str) -> tuple[Optional[Dict[str, Any]], Optional[str]]:
    """Parse Excel file and extract data."""
    try:
        from openpyxl import load_workbook

        data = {
            "type": "excel",
            "sheets": [],
            "total_sheets": 0
        }

        wb = load_workbook(file_path, data_only=True)
        data["total_sheets"] = len(wb.sheetnames)

        # Parse first 3 sheets
        for sheet_name in wb.sheetnames[:3]:
            ws = wb[sheet_name]
            sheet_data = {
                "name": sheet_name,
                "rows": []
            }

            # Extract first 20 rows
            for row_idx, row in enumerate(ws.iter_rows(max_row=20, values_only=True), 1):
                sheet_data["rows"].append({
                    "row": row_idx,
                    "values": [str(v) if v is not None else "" for v in row]
                })

            data["sheets"].append(sheet_data)

        return data, None

    except ImportError:
        return None, "openpyxl library not installed"
    except Exception as e:
        return None, f"Excel parsing error: {str(e)}"


async def _parse_word(file_path: str) -> tuple[Optional[Dict[str, Any]], Optional[str]]:
    """Parse Word document and extract text."""
    try:
        from docx import Document

        data = {
            "type": "word",
            "paragraphs": [],
            "tables": [],
            "full_text": ""
        }

        doc = Document(file_path)

        # Extract paragraphs
        for para_idx, para in enumerate(doc.paragraphs):
            if para.text.strip():
                data["paragraphs"].append({
                    "index": para_idx,
                    "text": para.text[:500]
                })
                data["full_text"] += para.text + "\n"

        # Extract tables
        for table_idx, table in enumerate(doc.tables[:5]):  # First 5 tables
            table_data = {
                "index": table_idx,
                "rows": []
            }
            for row in table.rows[:10]:  # First 10 rows
                table_data["rows"].append([
                    cell.text[:100] for cell in row.cells
                ])
            data["tables"].append(table_data)

        # Truncate full text
        data["full_text"] = data["full_text"][:5000]
        return data, None

    except ImportError:
        return None, "python-docx library not installed"
    except Exception as e:
        return None, f"Word parsing error: {str(e)}"


async def _parse_text(file_path: str) -> tuple[Optional[Dict[str, Any]], Optional[str]]:
    """Parse plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()

        data = {
            "type": "text",
            "content": content[:5000],
            "total_length": len(content)
        }

        return data, None

    except Exception as e:
        return None, f"Text parsing error: {str(e)}"
